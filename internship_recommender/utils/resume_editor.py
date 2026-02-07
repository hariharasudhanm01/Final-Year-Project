"""
AI-Powered Resume Editor
Intelligently modifies resumes to match job descriptions while preserving layout
Uses Ollama for LLM-based text generation and python-docx for DOCX manipulation
"""
import re
import os
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
import json

try:
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

try:
    import requests
    REQUESTS_AVAILABLE = True
except ImportError:
    REQUESTS_AVAILABLE = False


@dataclass
class ResumeSection:
    """Represents a section of the resume"""
    name: str
    content: str
    start_index: int
    end_index: int


@dataclass
class ModificationSuggestion:
    """Represents a suggested modification"""
    section: str
    original_text: str
    suggested_text: str
    reason: str
    skill_added: Optional[str] = None
    impact_score: float = 0.0  # Expected ATS score increase


class ResumeEditor:
    """
    AI-powered resume editor using Ollama
    """
    
    def __init__(self, ollama_url: str = "http://localhost:11434", model: str = "tinyllama"):
        """
        Initialize the resume editor
        
        Args:
            ollama_url: URL of Ollama API
            model: Model to use (llama2, mistral, etc.)
        """
        self.ollama_url = ollama_url
        self.model = model
        self.ollama_available = self._check_ollama()
    
    def _check_ollama(self) -> bool:
        """Check if Ollama is running"""
        if not REQUESTS_AVAILABLE:
            print("Warning: requests library not available")
            return False
        
        try:
            response = requests.get(f"{self.ollama_url}/api/tags", timeout=2)
            return response.status_code == 200
        except Exception as e:
            print(f"Warning: Ollama not available at {self.ollama_url}: {e}")
            return False
    
    def create_resume_from_template(self, data: Dict[str, Any], template_name: str, output_path: str) -> bool:
        """
        Creates a new resume DOCX file from structured data using a specific template style.
        """
        if not DOCX_AVAILABLE:
            print("Error: python-docx not installed")
            return False
            
        try:
            doc = Document()
            
            # --- STYLES ---
            style = doc.styles['Normal']
            font = style.font
            
            if template_name.lower() == "modern":
                font.name = 'Calibri'
                font.size = Pt(11)
                header_color = RGBColor(0, 51, 102) # Dark Blue
            else: # Professional / Classic
                font.name = 'Times New Roman'
                font.size = Pt(12)
                header_color = RGBColor(0, 0, 0) # Black
                
            # --- HEADER (Name & Contact) ---
            p_name = doc.add_paragraph()
            p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            p_info = data.get('personal_info')
            if not isinstance(p_info, dict): p_info = {}
            
            run_name = p_name.add_run(p_info.get('name', 'Your Name'))
            run_name.bold = True
            run_name.font.size = Pt(24)
            run_name.font.color.rgb = header_color
            
            contact_info = []
            if p_info.get('email'):
                contact_info.append(p_info['email'])
            if p_info.get('phone'):
                contact_info.append(p_info['phone'])
            if p_info.get('linkedin') and 'Not Found' not in p_info['linkedin']:
                contact_info.append("LinkedIn")
            if p_info.get('github') and 'Not Found' not in p_info['github']:
                contact_info.append("GitHub")

            p_contact = doc.add_paragraph(" | ".join(contact_info))
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_contact.style = 'No Spacing'
            
            # Draw line
            if template_name.lower() == "modern":
                doc.add_paragraph().add_run().add_break()
            else:
                p_line = doc.add_paragraph()
                p_line.paragraph_format.space_after = Pt(10)
                # p_line.paragraph_format.check_box_borders = True 
            
            # --- HELPER FOR SECTIONS ---
            def add_section_header(title):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                r = p.add_run(title.upper())
                r.bold = True
                r.font.size = Pt(14)
                r.font.color.rgb = header_color
                
            # --- SUMMARY ---
            summary = p_info.get('summary')
            if summary:
                add_section_header("Professional Summary")
                doc.add_paragraph(summary)
            
            # --- SKILLS ---
            skills = data.get('skills', {})
            all_skills = []
            if isinstance(skills, list):
                all_skills = skills
            elif isinstance(skills, dict):
                for k, v in skills.items():
                    if isinstance(v, list): all_skills.extend(v)
            
            if all_skills:
                add_section_header("Skills")
                doc.add_paragraph(", ".join(all_skills))
                
            # --- EXPERIENCE ---
            experience = data.get('experience', [])
            if not isinstance(experience, list) and experience: experience = [experience] # safe fallback
            
            if experience and isinstance(experience, list):
                add_section_header("Experience")
                for job in experience:
                    if not isinstance(job, dict): continue
                    
                    p_job = doc.add_paragraph()
                    p_job.paragraph_format.space_after = Pt(0)
                    
                    # Role & Company -> "Software Engineer | Google"
                    r_role = p_job.add_run(f"{job.get('role', 'Role')} | {job.get('company', 'Company')}")
                    r_role.bold = True
                    
                    # Dates -> Right aligned usually, but simple appending for now
                    if job.get('duration'):
                        p_job.add_run(f"  ({job.get('duration')})").italic = True
                    
                    # Description bullets
                    desc = job.get('description', [])
                    if isinstance(desc, list):
                        for point in desc:
                            doc.add_paragraph(point, style='List Bullet')
                    else:
                        doc.add_paragraph(str(desc))
                        
            # --- PROJECTS ---
            projects = data.get('projects', [])
            if not isinstance(projects, list) and projects: projects = [projects]

            if projects and isinstance(projects, list):
                add_section_header("Projects")
                for proj in projects:
                    if not isinstance(proj, dict): continue

                    p_proj = doc.add_paragraph()
                    p_proj.paragraph_format.space_after = Pt(0)
                    
                    title = proj.get('title', 'Project')
                    tech = proj.get('technologies', '')
                    
                    r_title = p_proj.add_run(title)
                    r_title.bold = True
                    if tech:
                        p_proj.add_run(f" - {tech}").italic = True
                        
                    desc = proj.get('description', [])
                    if isinstance(desc, list):
                        for point in desc:
                            doc.add_paragraph(point, style='List Bullet')
                    else:
                        doc.add_paragraph(str(desc))

            # --- EDUCATION ---
            education = data.get('education', [])
            if isinstance(education, dict): education = [education] 
            
            if education and isinstance(education, list):
                add_section_header("Education")
                for edu in education:
                    if not isinstance(edu, dict): continue
                    
                    p_edu = doc.add_paragraph()
                    line = f"{edu.get('degree', 'Degree')} - {edu.get('institution', 'College')}"
                    if edu.get('year'):
                        line += f" ({edu.get('year')})"
                    p_edu.add_run(line)
                    if edu.get('score'):
                         doc.add_paragraph(f"Score: {edu.get('score')}")
            
            # --- CERTIFICATIONS ---
            certs = data.get('certifications', [])
            if isinstance(certs, list) and certs:
                add_section_header("Certifications")
                for cert in certs:
                    doc.add_paragraph(str(cert), style='List Bullet')
            
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error creating resume from template: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _call_ollama(self, prompt: str, system_prompt: str = None) -> str:
        """
        Call Ollama API to generate text
        
        Args:
            prompt: User prompt
            system_prompt: System/instruction prompt
            
        Returns:
            Generated text
        """
        if not self.ollama_available:
            return "[Ollama not available - please start Ollama service]"
        
        try:
            payload = {
                "model": self.model,
                "prompt": prompt,
                "stream": False
            }
            
            if system_prompt:
                payload["system"] = system_prompt
            
            response = requests.post(
                f"{self.ollama_url}/api/generate",
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                return result.get("response", "")
            else:
                return f"[Ollama error: {response.status_code}]"
        
        except Exception as e:
            print(f"Error calling Ollama: {e}")
            return f"[Error: {str(e)}]"
    
    def extract_sections(self, resume_text: str) -> List[ResumeSection]:
        """
        Extract sections from resume text
        
        Common sections: Summary, Experience, Education, Skills, Projects, Certifications
        """
        sections = []
        
        # Common section headers (case-insensitive)
        section_patterns = [
            r'(SUMMARY|PROFILE|OBJECTIVE)',
            r'(EXPERIENCE|WORK HISTORY|EMPLOYMENT)',
            r'(EDUCATION|ACADEMIC)',
            r'(SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES)',
            r'(PROJECTS|KEY PROJECTS)',
            r'(CERTIFICATIONS|CERTIFICATES)',
            r'(ACHIEVEMENTS|ACCOMPLISHMENTS)',
        ]
        
        lines = resume_text.split('\n')
        current_section = None
        current_content = []
        current_start = 0
        
        for i, line in enumerate(lines):
            line_upper = line.strip().upper()
            
            # Check if line is a section header
            is_header = False
            for pattern in section_patterns:
                if re.match(pattern, line_upper):
                    # Save previous section
                    if current_section:
                        sections.append(ResumeSection(
                            name=current_section,
                            content='\n'.join(current_content),
                            start_index=current_start,
                            end_index=i
                        ))
                    
                    current_section = line.strip()
                    current_content = []
                    current_start = i
                    is_header = True
                    break
            
            if not is_header and current_section:
                current_content.append(line)
        
        # Add last section
        if current_section:
            sections.append(ResumeSection(
                name=current_section,
                content='\n'.join(current_content),
                start_index=current_start,
                end_index=len(lines)
            ))
        
        return sections
    
    def suggest_resume_modifications(
        self,
        resume_text: str,
        jd_text: str,
        missing_skills: List[str],
        weak_skills: List[str],
        ats_score: float
    ) -> List[ModificationSuggestion]:
        """
        Generate intelligent modification suggestions using Ollama
        
        Args:
            resume_text: Original resume text
            jd_text: Job description text
            missing_skills: Skills missing from resume
            weak_skills: Skills that are weak matches
            ats_score: Current ATS score
            
        Returns:
            List of modification suggestions
        """
        suggestions = []
        
        # Extract sections
        sections = self.extract_sections(resume_text)
        
        # 1. Skills Section Enhancement
        skills_section = next((s for s in sections if 'SKILL' in s.name.upper()), None)
        if skills_section and missing_skills:
            suggestion = self._enhance_skills_section(
                skills_section.content,
                missing_skills,
                jd_text
            )
            if suggestion:
                suggestions.append(suggestion)
        
        # 2. Experience Section Enhancement
        exp_section = next((s for s in sections if 'EXPERIENCE' in s.name.upper()), None)
        if exp_section and (missing_skills or weak_skills):
            suggestion = self._enhance_experience_section(
                exp_section.content,
                missing_skills + weak_skills,
                jd_text
            )
            if suggestion:
                suggestions.append(suggestion)
        
        # 3. Summary/Objective Enhancement
        summary_section = next((s for s in sections if any(kw in s.name.upper() for kw in ['SUMMARY', 'PROFILE', 'OBJECTIVE'])), None)
        if summary_section:
            suggestion = self._enhance_summary_section(
                summary_section.content,
                jd_text,
                missing_skills
            )
            if suggestion:
                suggestions.append(suggestion)
        
        return suggestions
    
    def _enhance_skills_section(
        self,
        skills_content: str,
        missing_skills: List[str],
        jd_text: str
    ) -> Optional[ModificationSuggestion]:
        """Enhance skills section by adding missing skills"""
        
        if not missing_skills:
            return None
        
        # Generate prompt for Ollama
        system_prompt = """You are an expert professional resume writer. 
        Your task is to enhance the Skills section by naturally integrating missing skills. 
        Keep the format consistent with the original.
        CRITICAL: Your output MUST be in ENGLISH. Do not use any other language."""
        
        prompt = f"""Original Skills Section:
{skills_content}

Missing Skills to Add: {', '.join(missing_skills[:5])}

Job Description Context:
{jd_text[:500]}

Rewrite the Skills section to include the missing skills naturally. Maintain the original format and style.
Only output the enhanced Skills section, nothing else. Ensure all text is in English."""

        enhanced_text = self._call_ollama(prompt, system_prompt)
        
        if enhanced_text and not enhanced_text.startswith('['):
            return ModificationSuggestion(
                section="Skills",
                original_text=skills_content,
                suggested_text=enhanced_text.strip(),
                reason=f"Added {len(missing_skills[:5])} missing skills: {', '.join(missing_skills[:5])}",
                skill_added=', '.join(missing_skills[:5]),
                impact_score=len(missing_skills[:5]) * 3.0
            )
        
        return None
    
    def _enhance_experience_section(
        self,
        experience_content: str,
        target_skills: List[str],
        jd_text: str
    ) -> Optional[ModificationSuggestion]:
        """Enhance experience section to highlight target skills"""
        
        if not target_skills:
            return None
        
        system_prompt = """You are an expert professional resume writer. Enhance the Experience section by 
        rephrasing bullet points to highlight relevant skills from the job description. 
        Keep it truthful and maintain the original achievements.
        CRITICAL: Your output MUST be in ENGLISH. Do not use any other language."""
        
        prompt = f"""Original Experience Section:
{experience_content[:1000]}

Skills to Highlight: {', '.join(target_skills[:5])}

Job Description Keywords:
{jd_text[:500]}

Rewrite the experience bullet points to better highlight these skills. Keep the same roles and timeframes.
Only enhance the descriptions to better match the job requirements. Output only the enhanced section.
Ensure the output is entirely in ENGLISH."""

        enhanced_text = self._call_ollama(prompt, system_prompt)
        
        if enhanced_text and not enhanced_text.startswith('['):
            return ModificationSuggestion(
                section="Experience",
                original_text=experience_content[:500] + "...",
                suggested_text=enhanced_text.strip()[:500] + "...",
                reason=f"Rephrased to highlight: {', '.join(target_skills[:3])}",
                impact_score=5.0
            )
        
        return None
    
    def _enhance_summary_section(
        self,
        summary_content: str,
        jd_text: str,
        missing_skills: List[str]
    ) -> Optional[ModificationSuggestion]:
        """Enhance summary/objective to align with job description"""
        
        system_prompt = """You are an expert professional resume writer. Rewrite the professional summary to 
        align with the job description while staying truthful to the candidate's background.
        CRITICAL: Your output MUST be in ENGLISH. Do not use any other language."""
        
        prompt = f"""Original Summary:
{summary_content}

Job Description:
{jd_text[:400]}

Key Skills to Mention: {', '.join(missing_skills[:3])}

Rewrite the summary to better align with the job requirements. Output only the enhanced summary, 2-3 sentences max.
Ensure the output is entirely in ENGLISH."""

        enhanced_text = self._call_ollama(prompt, system_prompt)
        
        if enhanced_text and not enhanced_text.startswith('['):
            return ModificationSuggestion(
                section="Summary",
                original_text=summary_content,
                suggested_text=enhanced_text.strip(),
                reason="Aligned summary with job requirements",
                impact_score=4.0
            )
        
        return None

    def rewrite_text_segment(self, text: str, style: str = "professional") -> List[str]:
        """
        Rewrite a text segment in different styles or improved clarity.
        Returns a list of 3 variations.
        """
        system_prompt = f"""You are an expert professional resume writer. 
        Your task is to rewrite the provided text to be more {style}, impactful, and concise.
        Use strong action verbs.
        Output exactly 3 distinct variations labeled 1., 2., and 3.
        Do not include any conversational text or introductions. Just the numbered list.
        CRITICAL: Your output MUST be in ENGLISH. Do not use any other language."""
        
        prompt = f"Rewrite this resume content:\n'{text}'\n\nEnsure the output is 100% in English."
        
        response = self._call_ollama(prompt, system_prompt)
        
        # Parse response into list
        variations = []
        if response:
            # Simple parsing for numbered list
            lines = response.split('\n')
            for line in lines:
                clean_line = re.sub(r'^\d+[\.)]\s*', '', line).strip()
                if clean_line and len(clean_line) > 5:
                    variations.append(clean_line)
                    
        # Fallback if parsing fails or valid response is just text
        if not variations and response and not response.startswith('['):
            variations = [response.strip()]
            
        return variations[:3]

    def apply_modifications_to_docx(
        self,
        docx_path: str,
        suggestions: List[ModificationSuggestion],
        output_path: str
    ) -> Tuple[bool, str]:
        """
        Apply modifications to a DOCX file
        
        Args:
            docx_path: Path to original DOCX
            suggestions: List of approved modifications
            output_path: Path to save modified DOCX
            
        Returns:
            Tuple(True if successful, Error message if failed)
        """
        if not DOCX_AVAILABLE:
            msg = "Error: python-docx not installed"
            print(msg)
            return False, msg
        
        try:
            # Load document
            if not os.path.exists(docx_path):
                 return False, f"Original file not found: {docx_path}"

            doc = Document(docx_path)
            
            # Apply modifications section by section
            for suggestion in suggestions:
                if not suggestion: continue
                self._apply_section_modification(doc, suggestion)
            
            # Save modified document
            doc.save(output_path)
            return True, ""
        
        except PermissionError:
            msg = "Permission denied: Is the file open in Word? Please close it and try again."
            print(f"Error applying modifications: {msg}")
            return False, msg
        except Exception as e:
            import traceback
            traceback.print_exc()
            msg = f"Error applying modifications: {e}"
            print(msg)
            return False, msg
    
    def _apply_section_modification(self, doc: Document, suggestion: ModificationSuggestion):
        """Apply a single modification to the document"""
        
        # Find paragraphs containing the section name
        section_found = False
        target_para_idx = -1
        
        for i, para in enumerate(doc.paragraphs):
            # Check if this paragraph is a section header
            # Heuristic: Upper case or Title case with specific keywords, short length
            text_upper = para.text.strip().upper()
            if suggestion.section.upper() in text_upper and len(text_upper) < 50:
                if any(kw in text_upper for kw in [suggestion.section.upper(), 'SUMMARY', 'EXPERIENCE', 'SKILLS', 'EDUCATION', 'PROJECTS']):
                     section_found = True
                     target_para_idx = i
                     break
        
        if section_found and target_para_idx != -1:
            # We found the header at target_para_idx
            # The content usually follows immediately
            
            # 1. Identify the range of paragraphs to remove/replace
            start_content_idx = target_para_idx + 1
            end_content_idx = start_content_idx
            
            while end_content_idx < len(doc.paragraphs):
                p = doc.paragraphs[end_content_idx]
                t = p.text.strip()
                # Stop if we hit next section header
                # Heuristic: Uppercase, bold (often), short
                if t and t.isupper() and len(t) < 40 and any(kw in t for kw in ['EXPERIENCE', 'EDUCATION', 'SKILLS', 'PROJECTS', 'CERTIFICATIONS', 'ACHIEVEMENTS', 'SUMMARY']):
                     break
                end_content_idx += 1
            
            # 2. Prepare new content paragraphs
            # Split suggested text by newlines to create separate paragraphs
            new_lines = suggestion.suggested_text.split('\n')
            
            # 3. Apply changes via "insert and delete" or "overwrite and delete rest"
            # Strategy: Overwrite the first few paragraphs with new lines, delete any excess old paragraphs,
            # or insert excess new paragraphs.
            
            # We need to be careful not to break the document structure. 
            # Safest approach for preserving style:
            # - Use the style of the first content paragraph for all new paragraphs (if available)
            # - Or default to 'Normal'
            
            # Get style reference from first content paragraph if it exists
            style_ref = None
            if start_content_idx < len(doc.paragraphs):
                style_ref = doc.paragraphs[start_content_idx].style
            
            current_idx = start_content_idx
            
            for line in new_lines:
                line = line.strip()
                if not line: continue
                
                if current_idx < end_content_idx:
                    # Overwrite existing paragraph (preserves some paragraph-level properties)
                    p = doc.paragraphs[current_idx]
                    p.text = line
                    # Re-apply style if we have a reference, to ensure consistency
                    if style_ref:
                        p.style = style_ref
                    current_idx += 1
                else:
                    # Access the previous paragraph to insert *after* it
                    # python-docx doesn't have easy 'insert_paragraph_at_index' without internal API usage
                    # But we can use insert_paragraph_before on the *next* section header (if exists) 
                    # or add_paragraph if at end.
                    
                    if end_content_idx < len(doc.paragraphs):
                        # Insert before the next section header
                        next_header_para = doc.paragraphs[end_content_idx]
                        new_p = next_header_para.insert_paragraph_before(line, style=style_ref)
                        # Adjustment: insert_paragraph_before adds it to the document list *conceptually*,
                        # but we need to track that the 'end_content_idx' effectively shifts or stays?
                        # Actually insert_paragraph_before modifies the DOM. 
                        # We don't need to increment current_idx logic for deletion, just insertion.
                    else:
                        # Append to end of doc
                        doc.add_paragraph(line, style=style_ref)
            
            # 4. If we have leftover old paragraphs, delete them (clear text)
            while current_idx < end_content_idx:
                # We can't easily delete paragraphs in python-docx without accessing obscure XML parent
                # Easiest way "preserving layout" is to just empty the text
                p = doc.paragraphs[current_idx]
                p.text = ""
                # Optional: p._element.getparent().remove(p._element) to actually delete
                # But emptying text is safer to avoid breaking structure if tables/etc are involved (though these are paras)
                
                # Let's try to actually delete to avoid unlimited whitespace
                try:
                    p._element.getparent().remove(p._element)
                except Exception:
                    p.text = "" # Fallback
                
                current_idx += 1
                
        else:
            # If section not found, append at end
            doc.add_paragraph(f"\n{suggestion.section}", style='Heading 2')
            for line in suggestion.suggested_text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())


# Global instance
resume_editor = ResumeEditor()
